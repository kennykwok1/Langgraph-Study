import os
import re
from typing import Optional
def print_log(level: str, message: str):
	print(f"[{level.upper()}] {message}")

def excel_to_markdown(file_path: str) -> str:
	try:
		import pandas as pd
		import openpyxl
	except ImportError as e:
		print_log('error', '需要安装 pandas 和 openpyxl 库来处理 Excel 文件。')
	md = []
	try:
		xls = pd.ExcelFile(file_path, engine='openpyxl')
		for sheet_name in xls.sheet_names:
			print_log('info', f'正在处理 Sheet: {sheet_name}')
			df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str, engine='openpyxl')
			md.append(f"\n## {sheet_name}\n")
			if df.empty:
				print_log('warning', f'Sheet {sheet_name} 为空。')
				continue
			# 转换为 Markdown 表格
			header = '| ' + ' | '.join(df.columns) + ' |'
			sep = '| ' + ' | '.join(['---'] * len(df.columns)) + ' |'
			md.append(header)
			md.append(sep)
			for _, row in df.iterrows():
				row_str = '| ' + ' | '.join(str(cell) if cell is not None else '' for cell in row) + ' |'
				md.append(row_str)
	except Exception as e:
		print_log('error', f'Excel 读取失败: {e}')
		raise e
	return '\n'.join(md)

def docx_to_markdown(file_path: str) -> str:
	try:
		from docx import Document
	except ImportError as e:
		print_log('error', '需要安装 python-docx 库来处理 Word 文件。')
		raise e
	md = []
	try:
		doc = Document(file_path)
		print_log('info', '正在处理 Word 文档段落和表格...')
		# 先处理段落
		for para in doc.paragraphs:
			text = ''
			for run in para.runs:
				# 检查 run 是否有超链接
				if run.hyperlink:
					text += f"[{run.text}]({run.hyperlink.target})"
				else:
					text += run.text
			text = convert_links_in_text(text)
			if text.strip():
				md.append(text)
		# 再处理表格
		for i, table in enumerate(doc.tables):
			md.append(f"\n### 表格 {i+1}\n")
			rows = list(table.rows)
			if not rows:
				print_log('warning', f'表格 {i+1} 为空。')
				continue
			# 表头
			header_cells = [cell.text.strip() for cell in rows[0].cells]
			header = '| ' + ' | '.join(header_cells) + ' |'
			sep = '| ' + ' | '.join(['---'] * len(header_cells)) + ' |'
			md.append(header)
			md.append(sep)
			for row in rows[1:]:
				row_cells = [cell.text.strip() for cell in row.cells]
				row_str = '| ' + ' | '.join(row_cells) + ' |'
				md.append(row_str)
	except Exception as e:
		print_log('error', f'Word 读取失败: {e}')
		raise e
	return '\n'.join(md)

def convert_links_in_text(text: str) -> str:
	# 将文本中的 http(s) 链接转为 markdown 链接
	url_pattern = re.compile(r'(https?://[\w\-./?%&=#:]+)')
	def repl(match):
		url = match.group(1)
		return f'[{url}]({url})'
	return url_pattern.sub(repl, text)

def file_to_markdown(input_path: str, output_path: Optional[str] = None) -> str:
	ext = os.path.splitext(input_path)[1].lower()
	if ext in ['.xlsx', '.xls']:
		print_log('info', f'检测到 Excel 文件: {input_path}')
		md = excel_to_markdown(input_path)
	elif ext == '.docx':
		print_log('info', f'检测到 Word 文件: {input_path}')
		md = docx_to_markdown(input_path)
	else:
		print_log('error', f'不支持的文件类型: {ext}')
		raise ValueError(f'不支持的文件类型: {ext}')

	# 处理输出路径
	if output_path is None:
		output_path = os.path.splitext(input_path)[0] + '.md'
	try:
		dir_name = os.path.dirname(output_path)
		if dir_name and not os.path.exists(dir_name):
			os.makedirs(dir_name, exist_ok=True)
			print_log('info', f'已创建目录: {dir_name}')
		with open(output_path, 'w', encoding='utf-8') as f:
			f.write(md)
		print_log('info', f'已写入文件: {output_path}')
	except Exception as e:
		print_log('error', f'写入 Markdown 文件失败: {e}')
		raise e
	return md


if __name__ == "__main__":
	import argparse
	parser = argparse.ArgumentParser(description="将 Excel 或 Word 文档内容转换为 Markdown 格式")
	parser.add_argument('--source', '-s', required=True, help='源文档路径（Excel 或 Word 文件）')
	parser.add_argument('--target', '-t', required=False, help='目标 Markdown 文件路径（可选）')
	args = parser.parse_args()
	try:
		md_str = file_to_markdown(args.source, args.target)
		print_log('info', '转换完成。')
	except Exception as e:
		print_log('error', f'转换失败: {e}')
